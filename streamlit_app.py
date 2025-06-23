
import streamlit as st
import pandas as pd
import io
import fitz
import re
from docx import Document

st.set_page_config(page_title="Calculadora EPH – Informe Automático", layout="wide")
st.title("📊 Calculadora EPH – Informe Excel + Word")

anio = st.selectbox("📅 Seleccioná el año de la base", ["2017", "2018", "2019", "2020", "2021", "2022", "2023"])
hogares_file = st.file_uploader("🏠 Base de Hogares anual (.xlsx)", type="xlsx")
individuos_file = st.file_uploader("👤 Base de Individuos anual (.xlsx)", type="xlsx")
instructivo_pdf = st.file_uploader("📄 Instructivo PDF", type="pdf")

def limpiar_descripcion_variable(desc):
    desc = desc.replace(".....", "").replace("....", "").replace("...", "").strip()
    correcciones = {
        "Tiene agua": "Acceso al agua",
        "El agua es de": "Fuente de agua",
        "¿tiene baño/letrina?": "Tiene baño o letrina",
        "El baño o letrina está": "Ubicación del baño o letrina",
        "El baño tiene": "Tipo de baño",
        "El desague del baño es": "Desagüe del baño",
        "La vivienda está ubicada cerca de basural/es(3": "Proximidad a basural",
        "La vivienda está ubicada en zona inundable": "Zona inundable",
        "La vivienda está ubicada en villa de emergencia": "Vivienda en villa de emergencia"
    }
    for parcial, reemplazo in correcciones.items():
        if parcial.lower() in desc.lower():
            return reemplazo
    return desc.strip().capitalize()

def extraer_diccionario_desde_pdf(pdf_file):
    text = ""
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page in doc:
        text += page.get_text()
    doc.close()
    regex = re.compile(r"^(\w{2,})\s+[NC]\(\d+\)\s+(.+)$", re.MULTILINE)
    matches = regex.findall(text)
    return {codigo.strip(): limpiar_descripcion_variable(desc) for codigo, desc in matches}

def generar_informe_word(anio):
    doc = Document()
    doc.add_heading(f"Informe Interpretativo EPH – Anual {anio}", level=1)
    doc.add_heading("🏠 Base de Hogares – Interpretación", level=2)
    doc.add_paragraph("El análisis incluye distribución regional, condiciones habitacionales, acceso a servicios básicos y tipología de vivienda.")
    doc.add_heading("👤 Base de Individuos – Interpretación", level=2)
    doc.add_paragraph("Se analiza la distribución por sexo, edad, nivel educativo, condición de actividad e ingresos.")
    doc.add_heading("📌 Conclusión General", level=2)
    doc.add_paragraph("Este informe permite identificar patrones sociales y económicos de la población urbana argentina para el año seleccionado.")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if hogares_file and individuos_file and instructivo_pdf:
    mapa = extraer_diccionario_desde_pdf(instructivo_pdf)
    df_hogar = pd.read_excel(hogares_file)
    df_ind = pd.read_excel(individuos_file)

    if mapa:
        df_hogar = df_hogar.rename(columns=mapa)
        df_ind = df_ind.rename(columns=mapa)

    if "CODUSU" in df_hogar.columns and "NRO_HOGAR" in df_hogar.columns:
        df_hogar = df_hogar.drop_duplicates(subset=["CODUSU", "NRO_HOGAR"])
    if all(x in df_ind.columns for x in ["CODUSU", "NRO_HOGAR", "COMPONENTE"]):
        df_ind = df_ind.drop_duplicates(subset=["CODUSU", "NRO_HOGAR", "COMPONENTE"])

    # Usar columnas nominales o crudas
    posibles_hogar = ["ingreso", "región", "agua", "baño", "vivienda", "ipcf", "itf", "PONDIH"]
    posibles_ind = ["sexo", "edad", "educ", "actividad", "ingreso", "ESTADO", "CH04", "CH05", "NIVEL_ED", "ITF", "IPCF"]

    cols_hogar = [c for c in df_hogar.columns if any(x in c.lower() for x in posibles_hogar)]
    cols_ind = [c for c in df_ind.columns if any(x in c.lower() for x in posibles_ind)]

    if not cols_hogar:
        st.warning("No se encontraron columnas clave en la base de hogares. Se usarán columnas originales si son útiles.")
        cols_hogar = df_hogar.columns[:10].tolist()

    if not cols_ind:
        st.warning("No se encontraron columnas clave en la base de individuos. Se usarán columnas originales si son útiles.")
        cols_ind = df_ind.columns[:10].tolist()

    resumen_hogar = df_hogar[cols_hogar].describe(include="all").transpose()
    resumen_ind = df_ind[cols_ind].describe(include="all").transpose()

    output_excel = io.BytesIO()
    with pd.ExcelWriter(output_excel, engine="openpyxl") as writer:
        resumen_hogar.to_excel(writer, sheet_name="Resumen Hogares")
        resumen_ind.to_excel(writer, sheet_name="Resumen Individuos")
    output_excel.seek(0)

    output_word = generar_informe_word(anio)

    st.success("✅ Análisis generado.")
    st.download_button("📥 Descargar Excel", data=output_excel, file_name=f"informe_eph_{anio}.xlsx")
    st.download_button("📥 Descargar Informe Interpretativo (Word)", data=output_word, file_name=f"informe_eph_{anio}.docx")
else:
    st.info("📥 Subí las bases de hogares, individuos y el instructivo PDF para comenzar.")
